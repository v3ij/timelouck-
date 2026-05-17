import docx

def docx_to_txt(docx_path, txt_path):
    doc = docx.Document(docx_path)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    
    # Also get tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            # De-duplicate adjacent identical cells due to merged cells
            cleaned_row = []
            for item in row_text:
                if not cleaned_row or cleaned_row[-1] != item:
                    cleaned_row.append(item)
            fullText.append(" | ".join(cleaned_row))
            
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(fullText))

docx_to_txt(
    "c:\\Users\\ayoub\\OneDrive\\سطح المكتب\\africano project\\TimeWallet_UI_UX_Specification_Document.docx",
    "c:\\Users\\ayoub\\OneDrive\\سطح المكتب\\africano project\\scratch\\TimeWallet_UI_UX_Specification_Document.txt"
)
print("Docx conversion done!")
